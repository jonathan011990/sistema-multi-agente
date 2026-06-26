import os
from flask import Flask, render_template, request, jsonify, send_file
from groq import Groq
from ddgs import DDGS
from docx import Document
from concurrent.futures import ThreadPoolExecutor
import io




app = Flask(__name__)
client = Groq(api_key="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")

def buscar_web(pregunta, max_results=5):
    with DDGS() as ddgs:
        resultados = list(ddgs.text(pregunta, max_results=max_results))
    return "\n".join([r["body"] for r in resultados])



def agente_investigador(tema):
    print("🔍 Agente Investigador trabajando...")
    
    client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
    # Las 3 búsquedas en paralelo
    with ThreadPoolExecutor(max_workers=3) as executor:
        f1 = executor.submit(buscar_web, f"¿Qué es {tema}?")
        f2 = executor.submit(buscar_web, f"{tema} últimas tendencias 2025")
        f3 = executor.submit(buscar_web, f"{tema} aplicaciones prácticas ejemplos")
        info1 = f1.result()
        info2 = f2.result()
        info3 = f3.result()

    resumen = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        max_tokens=4000,
        messages=[
            {"role": "system", "content": "Eres un investigador. Resume y organiza la información en puntos clave."},
            {"role": "user", "content": f"Información 1:\n{info1}\n\nInformación 2:\n{info2}\n\nInformación 3:\n{info3}\n\nOrganiza todo sobre: {tema}"}
        ]
    )
    return resumen.choices[0].message.content

def agente_redactor(tema, investigacion):

    client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

    informe = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        max_tokens=4000,
        messages=[
            {"role": "system", "content": """Eres un redactor profesional. Genera un informe completo con:
- Introducción
- Puntos principales
- Aplicaciones prácticas
- Conclusión"""},
            {"role": "user", "content": f"Tema: {tema}\n\nInvestigación:\n{investigacion}"}
        ]
    )
    return informe.choices[0].message.content

def agente_corrector(informe):

    client = Groq(api_key=os.environ.get("GROQ_API_KEY"))

    correccion = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        max_tokens=4000,
        messages=[
            {"role": "system", "content": """Eres un editor profesional. Revisa y mejora:
- Ortografía y gramática
- Claridad y coherencia
- Estructura y fluidez
Devuelve el informe completo corregido."""},
            {"role": "user", "content": f"Corrige y mejora este informe:\n\n{informe}"}
        ]
    )
    return correccion.choices[0].message.content


def agente_traductor(informe):

    client = Groq(api_key=os.environ.get("GROQ_API_KEY"))
    
    print("🌐 Agente Traductor trabajando...")
    traduccion = client.chat.completions.create(
        model="llama-3.1-8b-instant",
        max_tokens=4000,
        messages=[
            {"role": "system", "content": "Eres un traductor profesional. Traduce el siguiente informe al inglés manteniendo el formato y estructura original."},
            {"role": "user", "content": f"Traduce este informe al inglés:\n\n{informe}"}
        ]
    )
    return traduccion.choices[0].message.content


def guardar_word(tema, informe):
    doc = Document()
    doc.add_heading(f"Informe: {tema}", 0)
    for linea in informe.split("\n"):
        if linea.startswith("#"):
            doc.add_heading(linea.replace("#", "").strip(), level=1)
        elif linea.strip():
            doc.add_paragraph(linea)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/generar", methods=["POST"])
def generar():
    tema = request.json.get("tema")
    investigacion = agente_investigador(tema)
    informe = agente_redactor(tema, investigacion)
    informe_corregido = agente_corrector(informe)
    return jsonify({"informe": informe_corregido})

@app.route("/descargar", methods=["POST"])
def descargar():
    tema = request.json.get("tema")
    informe = request.json.get("informe")
    buffer = guardar_word(tema, informe)
    return send_file(buffer, as_attachment=True, download_name=f"informe_{tema}.docx", mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

@app.route("/traducir", methods=["POST"])
def traducir():
    informe = request.json.get("informe")
    traduccion = agente_traductor(informe)
    return jsonify({"traduccion": traduccion})

if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
    app.run(debug=True)